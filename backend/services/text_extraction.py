import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path: Path) -> str:

    suffix = file_path.suffix.lower()

    if suffix in (".txt", ".md"):
        return _extract_plaintext(file_path)
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)

    raise ValueError(
        f"Unsupported file type '{suffix}'. "
        "Supported: .txt, .md, .pdf, .docx"
    )



def _extract_plaintext(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        raise ValueError(f"File '{path.name}' is empty.")
    return text




def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("pypdf is not installed. Run: uv add pypdf") from e

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise ValueError(f"Could not open PDF '{path.name}': {e}") from e

    if reader.is_encrypted:
        raise ValueError(
            f"PDF '{path.name}' is password-protected and cannot be processed."
        )

    pages_text: list[str] = []
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            page_text = (page.extract_text() or "").strip()
            if page_text:
                pages_text.append(page_text)
        except Exception as e:
            logger.warning("PDF page %d extraction error in '%s': %s", page_num, path.name, e)

    if not pages_text:
        raise ValueError(
            f"PDF '{path.name}' contains no extractable text. "
            "It may be a scanned image-only PDF — convert it to a searchable PDF first."
        )

    full_text = "\n\n".join(pages_text)
    logger.info("PDF '%s': extracted %d chars from %d pages", path.name, len(full_text), len(pages_text))
    return full_text




def _extract_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise RuntimeError("python-docx is not installed. Run: uv add python-docx") from e

    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        raise ValueError(f"Could not open DOCX '{path.name}': {e}") from e

    parts: list[str] = []

    # Paragraphs (headings, body text, list items)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables — every cell, joined with a pipe separator per row
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))

    if not parts:
        raise ValueError(f"DOCX '{path.name}' contains no extractable text.")

    full_text = "\n\n".join(parts)
    logger.info("DOCX '%s': extracted %d chars from %d text blocks", path.name, len(full_text), len(parts))
    return full_text
